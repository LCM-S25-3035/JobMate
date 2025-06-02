import streamlit as st
import json
import os
from docx import Document
import google.generativeai as genai
from dotenv import load_dotenv
import re

load_dotenv()
genai.configure(api_key=os.getenv("GEMINI_API_KEY"))
model_gemini = "models/gemini-2.0-flash"

model = genai.GenerativeModel(
    model_gemini,
    system_instruction="You are a helpful assistant writing personalized, ATS-optimized cover letters.",
    generation_config={
        "temperature": 0.2,
        "top_p": 1.0,
        "top_k": 40,
        "max_output_tokens": 1024
    }
)

st.set_page_config(page_title="AI Cover Letter Generator", layout="centered")
st.title("📄 AI Cover Letter Generator")

st.markdown("Follow the steps below to generate a customized, ATS-friendly cover letter:")

# Step 1: Paste Job Description
st.subheader("Step 1: Paste Job Posting (JSON)")
job_input = st.text_area("Paste the full job posting JSON here", height=300)

# Step 2: Paste Resume
st.subheader("Step 2: Paste Resume (JSON)")
resume_input = st.text_area("Paste the full resume JSON here", height=300)

# Step 3: Generate
st.subheader("Step 3: Generate Cover Letter")
if st.button("🚀 Generate Now"):
    try:
        job = json.loads(job_input)
        resume = json.loads(resume_input)

        # Debug display + fail-safe fallback parsing
        st.subheader("✅ Resume Structure Debug")
        st.json(resume)

        if "personal_information" not in resume:
            st.warning("⚠️ 'personal_information' field is missing. Using default values.")
        if "education" not in resume:
            st.warning("⚠️ 'education' field is missing.")
        if "soft_skills" not in resume or "technical_skills" not in resume:
            st.warning("⚠️ Skills section is incomplete.")

        # Safe fallback extraction
        name = resume.get("personal_information", {}).get("name", "Your Name")
        email = resume.get("personal_information", {}).get("email", "you@example.com")
        education = ", ".join([edu.get("degree", "Unknown Degree") for edu in resume.get("education", [])])
        skills = ", ".join(resume.get("soft_skills", []) + resume.get("technical_skills", []))

        # Load template
        template_path = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "template", "cover_letter_template1.docx"))
        # Debug
        print("📁 Template path resolved to:", template_path)
        st.code(template_path, language="text")

        doc = Document(template_path)
        template_text = "\n".join(p.text for p in doc.paragraphs)

        # Gemini prompt
        prompt = f"""
        You're an expert resume assistant. Using the following resume and job description, generate a custom cover letter that is professional, ATS-optimized, and easy to read.

        ✅ Structure:
        - Begin with a formal greeting.
        - Paragraph 1: Mention the role and where it was found. State interest and motivation.
        - Paragraph 2: Show 2–3 experiences that directly align with job responsibilities.
        - Paragraph 3: Bullet 3 core skills or strengths.
        - Final paragraph: Invite for interview + appreciation.
        - Sign off with name.

        ✅ Formatting:
        - Use clear paragraph breaks (two line breaks between paragraphs).
        - Use bullet points for key strengths.

        Only use facts from the resume. Do not invent degrees, companies, or skills.

        === Resume Summary ===
        Name: {name}
        Email: {email}
        Education: {education}
        Professional Summary: {resume.get("professional_summary", "")}
        Skills: {skills}
        Key Highlights:
        - Launched Carro Taiwan MVP and drove 15% conversion uplift
        - Delivered AI-powered pricing and API integrations across regional markets
        - Managed UAT, sprint planning, and stakeholder alignment for 3+ product rollouts

        === Job Info ===
        Job Title: {job.get("job_title", "")}
        Company: {job.get("company", {}).get("name", "")}
        Responsibilities: {", ".join(job.get("responsibilities", [])[:5])}
        Requirements: {", ".join(job.get("requirements", [])[:5])}
        Description: {job.get("job_description", "")}

        Return only the filled-in final cover letter, well formatted.
        """


        # Gemini call
        with st.spinner("Generating with Gemini..."):
            response = model.generate_content(prompt)
            output_text = response.text.strip()

            if output_text.startswith("```"):
                output_text = re.sub(r"^```(text|txt)?", "", output_text)
                output_text = output_text.rstrip("```").strip()
            output_text = re.sub(r"[\x00-\x1F\x7F]", "", output_text)

            doc_output = Document()
            for line in output_text.split("\n"):
                doc_output.add_paragraph(line)

            output_path = "cover_letter/generated_cover_letter.docx"
            os.makedirs("cover_letter", exist_ok=True)
            doc_output.save(output_path)

        # Download button
        with open(output_path, "rb") as f:
            st.success("✅ Cover letter generated successfully!")
            st.download_button(
                label="📄 Download Cover Letter",
                data=f,
                file_name="generated_cover_letter.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

    except Exception as e:
        st.error(f"❌ Error: {str(e)}")
